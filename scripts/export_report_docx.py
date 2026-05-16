"""
Raporu Word (.docx) formatina cevirir.
Kullanim: pip install python-docx && python scripts/export_report_docx.py
"""
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "rapor" / "Proje_Raporu.md"
OUT = ROOT / "rapor" / "Proje_Raporu.docx"


def main():
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("python-docx gerekli: pip install python-docx")
        return

    text = MD.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for line in text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip() == "---":
            continue
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(OUT)
    print(f"Kaydedildi: {OUT}")


if __name__ == "__main__":
    main()
